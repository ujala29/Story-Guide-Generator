import sys
if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8")
import argparse
import os
import re
import tempfile
from pathlib import Path

import pypandoc
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
# from docx2pdf import convert


# ── Config ────────────────────────────────────────────────────

BASE_DIR    = Path(__file__).resolve().parent.parent.parent
OUTPUT_ROOT = BASE_DIR / "output"

PAGE_BREAK = '\n\n```{=openxml}\n<w:p><w:r><w:br w:type="page"/></w:r></w:p>\n```\n\n'

# Pages whose slugs start with these keywords are pulled to the front (in this order).
# Works for any dashboard — main_page, main, overview, summary always come first.
FIRST_PAGE_KEYWORDS = ["main_page", "main", "overview", "summary"]

VISUAL_PRIORITY = {
    "card": 1, "trend": 2, "line": 2,
    "bar": 3,  "column": 3,
    "table": 4, "matrix": 4,
    "donut": 5, "pie": 5,
}


LOGO_PATH = BASE_DIR / "input" / "innovaccer.png"
BLACK = RGBColor(0x00, 0x00, 0x00)


def insert_cover_page(docx_path: str, dashboard: str):
    display_name = dashboard.replace("-", " ").title()
    doc = Document(docx_path)
    body = doc.element.body
    existing = list(body)

    # --- build cover content (appended temporarily at end, then moved to front) ---

    # top spacer — pushes logo ~1/3 down the page
    p_top = doc.add_paragraph()
    p_top.paragraph_format.space_before = Pt(130)
    p_top.paragraph_format.space_after  = Pt(0)

    # logo
    if LOGO_PATH.exists():
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_logo.paragraph_format.space_before = Pt(0)
        p_logo.paragraph_format.space_after  = Pt(0)
        p_logo.add_run().add_picture(str(LOGO_PATH), width=Inches(2.2))

    # mid spacer — pushes title toward vertical center
    p_mid = doc.add_paragraph()
    p_mid.paragraph_format.space_before = Pt(90)
    p_mid.paragraph_format.space_after  = Pt(0)

    # title line 1: "<Dashboard Name> Dashboard"
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_before = Pt(0)
    p1.paragraph_format.space_after  = Pt(6)
    r1 = p1.add_run(f"{display_name} Dashboard")
    r1.font.name      = 'Calibri'
    r1.font.size      = Pt(22)
    r1.font.bold      = True
    r1.font.color.rgb = BLACK

    # title line 2: "Story Guide"
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after  = Pt(0)
    r2 = p2.add_run("Story Guide")
    r2.font.name      = 'Calibri'
    r2.font.size      = Pt(22)
    r2.font.bold      = True
    r2.font.color.rgb = BLACK

    # page break
    p_br = doc.add_paragraph()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    p_br.add_run()._r.append(br)

    # move all new elements to front of body (before first existing child)
    new_elems = [c for c in body if c not in existing]
    if not existing:
        raise RuntimeError("Pandoc produced an empty document body — check input markdown")
    ref = existing[0]
    for elem in new_elems:
        body.remove(elem)
        ref.addprevious(elem)

    doc.save(docx_path)
    print(f"  ✓ Cover page added: {display_name} Dashboard")


def update_footer_name(docx_path: str, dashboard: str):
    display_name = dashboard.replace("-", " ").title()
    doc = Document(docx_path)
    placeholder = "[Dashboard Name]"
    for section in doc.sections:
        for para in section.footer.paragraphs:
            # Word often splits text across multiple runs — check full para text first
            if placeholder not in para.text:
                continue
            # Merge all run text, replace, write back into the first run, clear the rest
            full_text = "".join(r.text for r in para.runs)
            new_text  = full_text.replace(placeholder, display_name)
            if para.runs:
                para.runs[0].text = new_text
                for run in para.runs[1:]:
                    run.text = ""
    doc.save(docx_path)
    print(f"  ✓ Footer updated: {display_name}")


def style_tables(docx_path: str):
    doc = Document(docx_path)
    for table in doc.tables:
        # full-width
        tbl = table._tbl
        tblPr = tbl.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        for old in tblPr.findall(qn('w:tblW')):
            tblPr.remove(old)
        tblW = OxmlElement('w:tblW')
        tblW.set(qn('w:w'), '5000')
        tblW.set(qn('w:type'), 'pct')
        tblPr.append(tblW)

        # auto-fit columns so no column gets crushed
        for old in tblPr.findall(qn('w:tblLayout')):
            tblPr.remove(old)
        tblLayout = OxmlElement('w:tblLayout')
        tblLayout.set(qn('w:type'), 'autofit')
        tblPr.append(tblLayout)

        for row_idx, row in enumerate(table.rows):
            is_header = row_idx == 0
            bg = '2E2E2E' if is_header else ('FFFFFF' if row_idx % 2 == 1 else 'F5F5F5')
            for cell in row.cells:
                tcPr = cell._tc.get_or_add_tcPr()
                # borders
                for old in tcPr.findall(qn('w:tcBorders')):
                    tcPr.remove(old)
                borders = OxmlElement('w:tcBorders')
                for edge in ('top', 'left', 'bottom', 'right'):
                    el = OxmlElement(f'w:{edge}')
                    el.set(qn('w:val'), 'single')
                    el.set(qn('w:sz'), '4')
                    el.set(qn('w:space'), '0')
                    el.set(qn('w:color'), 'CCCCCC')
                    borders.append(el)
                tcPr.append(borders)
                # background
                for old in tcPr.findall(qn('w:shd')):
                    tcPr.remove(old)
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), bg)
                tcPr.append(shd)
                # padding
                for old in tcPr.findall(qn('w:tcMar')):
                    tcPr.remove(old)
                tcMar = OxmlElement('w:tcMar')
                for side in ('top', 'left', 'bottom', 'right'):
                    m = OxmlElement(f'w:{side}')
                    m.set(qn('w:w'), str(6 * 20))
                    m.set(qn('w:type'), 'dxa')
                    tcMar.append(m)
                tcPr.append(tcMar)
                # font
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.name = 'Calibri'
                        run.font.size = Pt(10)
                        if is_header:
                            run.font.bold = True
                            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        else:
                            run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    doc.save(docx_path)
    print(f"  ✓ Table styling applied ({len(doc.tables)} tables)")


def insert_word_toc(docx_path: str):
    doc  = Document(docx_path)
    body = doc.element.body

    # ── "Table of Contents" heading ───────────────────────────
    toc_heading = OxmlElement('w:p')
    h_pPr   = OxmlElement('w:pPr')
    h_style = OxmlElement('w:pStyle')
    h_style.set(qn('w:val'), 'Heading1')
    h_pPr.append(h_style)
    toc_heading.append(h_pPr)
    h_r = OxmlElement('w:r')
    h_t = OxmlElement('w:t')
    h_t.text = 'Table of Contents'
    h_r.append(h_t)
    toc_heading.append(h_r)

    # ── TOC field: headings 1-3, hyperlinked, with page numbers ─
    toc_para = OxmlElement('w:p')

    r1 = OxmlElement('w:r')
    fc1 = OxmlElement('w:fldChar')
    fc1.set(qn('w:fldCharType'), 'begin')
    fc1.set(qn('w:dirty'), 'true')
    r1.append(fc1)
    toc_para.append(r1)

    r2 = OxmlElement('w:r')
    instr = OxmlElement('w:instrText')
    instr.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    r2.append(instr)
    toc_para.append(r2)

    r3 = OxmlElement('w:r')
    fc3 = OxmlElement('w:fldChar')
    fc3.set(qn('w:fldCharType'), 'separate')
    r3.append(fc3)
    toc_para.append(r3)

    r4 = OxmlElement('w:r')
    t4 = OxmlElement('w:t')
    t4.text = 'Right-click → Update Field to generate table of contents'
    r4.append(t4)
    toc_para.append(r4)

    r5 = OxmlElement('w:r')
    fc5 = OxmlElement('w:fldChar')
    fc5.set(qn('w:fldCharType'), 'end')
    r5.append(fc5)
    toc_para.append(r5)

    # ── Page break after TOC ───────────────────────────────────
    pb_para = OxmlElement('w:p')
    pb_r    = OxmlElement('w:r')
    pb_br   = OxmlElement('w:br')
    pb_br.set(qn('w:type'), 'page')
    pb_r.append(pb_br)
    pb_para.append(pb_r)

    # Insert at front — cover page will be prepended after this
    first = body[0]
    for elem in [toc_heading, toc_para, pb_para]:
        first.addprevious(elem)

    doc.save(docx_path)
    print('  ✓ TOC field inserted (right-click → Update Field in Word)')


def style_section_title(docx_path: str, title: str, size_pt: int = 16):
    doc = Document(docx_path)
    for para in doc.paragraphs:
        if para.text.strip() == title:
            para.style = doc.styles["Normal"]
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after  = Pt(6)
            for run in para.runs:
                run.font.name  = "Calibri"
                run.font.size  = Pt(size_pt)
                run.font.bold  = True
                run.font.color.rgb = BLACK
    doc.save(docx_path)


def format_faq_md(content: str) -> str:
    """
    Converts FAQ markdown (bold question + inline answer, no blank line between)
    into numbered Q&A blocks with question as ### heading and answer below it.

    Input:
        **Question?**
        Answer text.

    Output:
        ### Q1. Question?

        Answer text.
    """
    output = []
    q_num  = 0

    # Split into blocks by --- separator; strip each block
    raw_blocks = re.split(r'\n\s*---\s*\n', content)

    for block in raw_blocks:
        block = block.strip()
        if not block:
            continue

        # Section heading block — keep as-is
        if block.startswith('##'):
            output.append(block)
            output.append('\n---\n')
            continue

        # Q&A block: first line is **Question?**, rest is the answer
        lines      = block.split('\n', 1)
        q_line     = lines[0].strip()
        answer     = lines[1].strip() if len(lines) > 1 else ''

        if q_line.startswith('**') and q_line.endswith('**'):
            q_num += 1
            question = q_line.strip('*').strip()
            output.append(f'### Q{q_num}. {question}\n\n{answer}')
        else:
            # Not a Q&A block — keep as-is
            output.append(block)

    return '\n'.join(output).rstrip('\n') + '\n\n'


def read_file(path: Path) -> str:
    content = path.read_text(encoding="utf-8")
    content = re.sub(
        r'^\*\*Widget:\s*(.+?)\*\*\s*$',
        r'## \1',
        content,
        flags=re.MULTILINE,
    )
    return content.rstrip('\n') + '\n\n'


def sort_visuals(files):
    def priority(f):
        name = f.name.lower()
        for key, val in VISUAL_PRIORITY.items():
            if key in name:
                return val
        return 999
    return sorted(files, key=lambda f: (priority(f), f.name.lower()))


def sort_pages(page_dirs):
    def _priority(page_dir):
        slug = page_dir.name.lower()
        for i, kw in enumerate(FIRST_PAGE_KEYWORDS):
            if slug.startswith(kw):
                return i
        return len(FIRST_PAGE_KEYWORDS)
    return sorted(page_dirs, key=lambda p: (_priority(p), p.name.lower()))


MAX_METRIC_ROWS = 10


def build_metric_catalog_section(dashboard: str) -> str:
    """Read metric_catalog.md, return first MAX_METRIC_ROWS data rows + truncation note."""
    catalog_path = (
        BASE_DIR / "output" / "dashboards" / dashboard / "metric_dictionary" / "metric_catalog.md"
    )
    if not catalog_path.exists():
        print(f"  ⚠ not found: {catalog_path}")
        return ""

    lines = catalog_path.read_text(encoding="utf-8").splitlines()

    header = None
    separator = None
    current: list[str] = []
    data_rows: list[str] = []

    for line in lines:
        if not line.strip():
            continue
        if header is None:
            if line.strip().startswith("|"):
                header = line.strip()
            continue
        if separator is None:
            if re.match(r"^\s*\|[-| :]+\|\s*$", line):
                separator = line.strip()
            continue
        # group multi-physical-line cells into one logical row
        if line.strip().startswith("|"):
            if current:
                data_rows.append(" ".join(current))
            current = [line.strip()]
        else:
            if current:
                current.append(line.strip())

    if current:
        data_rows.append(" ".join(current))

    total = len(data_rows)
    preview = data_rows[:MAX_METRIC_ROWS]
    table_md = "\n".join(x for x in [header, separator] + preview if x)

    note = ""
    if total > MAX_METRIC_ROWS:
        note = (
            f"\n\n> **Note:** Showing {len(preview)} of {total} measures. "
            f"For complete DAX definitions of all measures, refer to the "
            f"**Metric Catalog Excel file** embedded in this document.\n"
        )

    print(f"  + Metric Dictionary ({len(preview)} of {total} measures)")
    return table_md + note + "\n\n"


# ── Build markdown ────────────────────────────────────────────

def build_combined_md(dashboard: str = "risk-dash") -> str:
    chunks = []
    dash_out = BASE_DIR / "output" / "dashboards" / dashboard
    dashboard_overview_dir = dash_out / "dashboard_overview"
    filter_section_dir     = dash_out / "filter_section"
    page_wise_dir          = dash_out / "page_wise"
    visual_wise_dir        = dash_out / "visual_wise"
    glossary_faq_dir       = dash_out / "glossary_faq"

    # TOC is inserted as a real Word field in post-processing (insert_word_toc)
    # — not as markdown, so Word auto-generates page numbers + sub-sections

    # 1. Dashboard Overview  (md file has its own heading — no injected title)
    p = dashboard_overview_dir / "dashboard_overview.md"
    if p.exists():
        chunks.append(read_file(p))
        print("  + Dashboard Overview")
    else:
        print(f"  ⚠ not found: {p}")

    # 2. Global Filters
    p = filter_section_dir / "global_filters.md"
    if p.exists():
        chunks.append(PAGE_BREAK)
        chunks.append("# Global Filters\n\n")
        chunks.append(read_file(p))
        print("  + Global Filters")
    else:
        print(f"  ⚠ not found: {p}")

    # 3. Page-Wise Story
    p = page_wise_dir / "page_wise_story.md"
    if p.exists():
        chunks.append(PAGE_BREAK)
        chunks.append(read_file(p))
        print("  + Page-Wise Story")
    else:
        print(f"  ⚠ not found: {p}")

    # 4. Visual-Wise  (card -> trend -> bar -> table -> donut, overview page first)
    story_root = visual_wise_dir / "story_guide"
    if story_root.exists():
        count = 0
        for page_dir in sort_pages([d for d in story_root.iterdir() if d.is_dir()]):
            chunks.append(PAGE_BREAK)
            chunks.append(f"# {page_dir.name.replace('_', ' ').title()}\n\n")
            for vf in sort_visuals(list(page_dir.glob("*.md"))):
                chunks.append(read_file(vf))
                chunks.append("\n\n---\n\n")
                count += 1
        print(f"  + Visual files : {count}")
    else:
        print(f"  ⚠ not found: {story_root}")


    # 5. Metric Dictionary (first 10 measures + truncation note)
    metric_md = build_metric_catalog_section(dashboard)
    if metric_md:
        chunks.append(PAGE_BREAK)
        chunks.append("# Metric Dictionary\n\n")
        chunks.append(metric_md)

    # 6. FAQ
    p = glossary_faq_dir / "faq.md"
    if p.exists():
        chunks.append(PAGE_BREAK)
        chunks.append("# FAQ\n\n")
        chunks.append(format_faq_md(p.read_text(encoding="utf-8")))
        print("  + FAQ")
    else:
        print(f"  ⚠ not found: {p}")

    # 7. Glossary
    p = glossary_faq_dir / "glossary.md"
    if p.exists():
        chunks.append(PAGE_BREAK)
        chunks.append("# Glossary\n\n")
        chunks.append(read_file(p))
        print("  + Glossary")
    else:
        print(f"  ⚠ not found: {p}")

    return "".join(chunks)


# ── Main ──────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description="Generate Story Guide Word document")
    parser.add_argument("--dashboard", default="risk-dash",
                        help="Dashboard name (e.g. risk-dash, pac-dash)")
    args = parser.parse_args()

    md = build_combined_md(dashboard=args.dashboard)
    output_path = str(OUTPUT_ROOT / f"{args.dashboard}_story_guide.docx")

    with tempfile.NamedTemporaryFile(
        mode='w', suffix='.md', encoding='utf-8', delete=False
    ) as tmp:
        tmp.write(md)
        tmp_path = tmp.name

    try:
        # Fail fast if the output file is locked (e.g. open in Word)
        if Path(output_path).exists():
            try:
                with open(output_path, "ab"):
                    pass
            except PermissionError:
                print(f"\n  ERROR: Cannot write to '{output_path}'")
                print(f"  The file is open in another application (e.g. Microsoft Word).")
                print(f"  Please close it and run again.")
                sys.exit(1)

        extra_args = ["--standalone", "--wrap=none"]
        ref_doc = OUTPUT_ROOT / "reference.docx"
        if ref_doc.exists():
            extra_args.append(f"--reference-doc={ref_doc}")

        print("  Converting...")
        pypandoc.convert_file(tmp_path, "docx", outputfile=output_path, extra_args=extra_args)
        print("  Styling tables...")
        style_tables(output_path)
        style_section_title(output_path, "Page Wise Narrative", size_pt=16)
        print("  Inserting TOC...")
        insert_word_toc(output_path)
        print("  Adding cover page...")
        insert_cover_page(output_path, args.dashboard)
        update_footer_name(output_path, args.dashboard)
        print(f"  ✅ Word: {output_path}")

        # pdf_path = output_path.replace(".docx", ".pdf")
        # print("  Converting to PDF...")
        # convert(output_path, pdf_path)
        # print(f"  ✅ PDF : {pdf_path}")
    finally:
        os.unlink(tmp_path)


if __name__ == "__main__":
    main()
