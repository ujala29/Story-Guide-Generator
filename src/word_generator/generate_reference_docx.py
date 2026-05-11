# import os
# from pathlib import Path
# from docx import Document
# from docx.shared import Pt, Inches, RGBColor, Cm
# from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
# from docx.enum.style import WD_STYLE_TYPE
# from docx.oxml.ns import qn
# from docx.oxml import OxmlElement

# OUTPUT_ROOT = Path(__file__).resolve().parent.parent/ "output"


# def set_cell_background(cell, color_hex):
#     tc = cell._tc
#     tcPr = tc.get_or_add_tcPr()
#     shd = OxmlElement('w:shd')
#     shd.set(qn('w:val'), 'clear')
#     shd.set(qn('w:color'), 'auto')
#     shd.set(qn('w:fill'), color_hex.lstrip('#'))
#     tcPr.append(shd)


# def set_cell_padding(cell, padding_pt):
#     tc = cell._tc
#     tcPr = tc.get_or_add_tcPr()
#     tcMar = OxmlElement('w:tcMar')
#     for side in ('top', 'left', 'bottom', 'right'):
#         margin = OxmlElement(f'w:{side}')
#         margin.set(qn('w:w'), str(int(padding_pt * 20)))
#         margin.set(qn('w:type'), 'dxa')
#         tcMar.append(margin)
#     tcPr.append(tcMar)


# def add_border(table, color='CCCCCC', sz='4'):
#     for row in table.rows:
#         for cell in row.cells:
#             tc = cell._tc
#             tcPr = tc.get_or_add_tcPr()
#             tcBorders = OxmlElement('w:tcBorders')
#             for edge in ('top', 'left', 'bottom', 'right'):
#                 el = OxmlElement(f'w:{edge}')
#                 el.set(qn('w:val'), 'single')
#                 el.set(qn('w:sz'), sz)
#                 el.set(qn('w:space'), '0')
#                 el.set(qn('w:color'), color)
#                 tcBorders.append(el)
#             tcPr.append(tcBorders)


# def add_page_number(paragraph):
#     run1 = paragraph.add_run()
#     fldChar1 = OxmlElement('w:fldChar')
#     fldChar1.set(qn('w:fldCharType'), 'begin')
#     run1._r.append(fldChar1)

#     run2 = paragraph.add_run()
#     instrText = OxmlElement('w:instrText')
#     instrText.text = 'PAGE'
#     instrText.set(qn('xml:space'), 'preserve')
#     run2._r.append(instrText)

#     run3 = paragraph.add_run()
#     fldChar2 = OxmlElement('w:fldChar')
#     fldChar2.set(qn('w:fldCharType'), 'end')
#     run3._r.append(fldChar2)

#     for run in (run1, run2, run3):
#         run.font.name = 'Calibri'
#         run.font.size = Pt(9)
#         run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


# def get_or_add_style_pPr(style):
#     element = style._element
#     pPr = element.find(qn('w:pPr'))
#     if pPr is None:
#         pPr = OxmlElement('w:pPr')
#         element.append(pPr)
#     return pPr


# def force_style_rPr_color(style, r, g, b):
#     """Force run color on a style's rPr, bypassing theme color inheritance."""
#     element = style._element
#     rPr = element.find(qn('w:rPr'))
#     if rPr is None:
#         rPr = OxmlElement('w:rPr')
#         element.append(rPr)
#     # Remove any existing color element
#     for existing in rPr.findall(qn('w:color')):
#         rPr.remove(existing)
#     color_el = OxmlElement('w:color')
#     color_el.set(qn('w:val'), f'{r:02X}{g:02X}{b:02X}')
#     rPr.append(color_el)
#     # Also clear theme color so it doesn't override
#     theme_color = rPr.find(qn('w:themeColor'))
#     if theme_color is not None:
#         rPr.remove(theme_color)


# def patch_table_grid_borders(doc, color='CCCCCC', sz='4'):
#     """Add border definitions to the Table Grid table style."""
#     try:
#         tg = doc.styles['Table Grid']
#     except KeyError:
#         return
#     tbl_el = tg._element
#     tblPr = tbl_el.find(qn('w:tblPr'))
#     if tblPr is None:
#         tblPr = OxmlElement('w:tblPr')
#         tbl_el.append(tblPr)
#     # Remove stale borders element if present
#     for old in tblPr.findall(qn('w:tblBorders')):
#         tblPr.remove(old)
#     tblBorders = OxmlElement('w:tblBorders')
#     for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
#         el = OxmlElement(f'w:{edge}')
#         el.set(qn('w:val'), 'single')
#         el.set(qn('w:sz'), sz)
#         el.set(qn('w:space'), '0')
#         el.set(qn('w:color'), color)
#         tblBorders.append(el)
#     tblPr.append(tblBorders)


# def create_reference_doc():
#     doc = Document()

#     # ── Page Setup (A4) ───────────────────────────────────────
#     section = doc.sections[0]
#     section.page_width = Cm(21.0)
#     section.page_height = Cm(29.7)
#     section.top_margin = Inches(1)
#     section.bottom_margin = Inches(1)
#     section.left_margin = Inches(1)
#     section.right_margin = Inches(1)

#     # ── Normal / Body Text ────────────────────────────────────
#     normal = doc.styles['Normal']
#     normal.font.name = 'Calibri'
#     normal.font.size = Pt(11)
#     normal.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
#     normal.paragraph_format.space_after = Pt(6)
#     normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
#     normal.paragraph_format.line_spacing = 1.15

#     # ── Heading 1 ─────────────────────────────────────────────
#     h1 = doc.styles['Heading 1']
#     h1.font.name = 'Calibri'
#     h1.font.size = Pt(16)
#     h1.font.bold = True
#     h1.paragraph_format.space_before = Pt(24)
#     h1.paragraph_format.space_after = Pt(12)
#     h1.paragraph_format.keep_with_next = True
#     force_style_rPr_color(h1, 0x00, 0x00, 0x00)
#     pPr_h1 = get_or_add_style_pPr(h1)
#     pBdr = OxmlElement('w:pBdr')
#     bottom = OxmlElement('w:bottom')
#     bottom.set(qn('w:val'), 'single')
#     bottom.set(qn('w:sz'), '6')
#     bottom.set(qn('w:space'), '1')
#     bottom.set(qn('w:color'), '000000')
#     pBdr.append(bottom)
#     pPr_h1.append(pBdr)

#     # ── Heading 2 ─────────────────────────────────────────────
#     h2 = doc.styles['Heading 2']
#     h2.font.name = 'Calibri'
#     h2.font.size = Pt(13)
#     h2.font.bold = True
#     h2.paragraph_format.space_before = Pt(12)
#     h2.paragraph_format.space_after = Pt(6)
#     force_style_rPr_color(h2, 0x00, 0x00, 0x00)

#     # ── Heading 3 ─────────────────────────────────────────────
#     h3 = doc.styles['Heading 3']
#     h3.font.name = 'Calibri'
#     h3.font.size = Pt(11)
#     h3.font.bold = True
#     h3.font.italic = True
#     h3.paragraph_format.space_before = Pt(8)
#     h3.paragraph_format.space_after = Pt(4)
#     force_style_rPr_color(h3, 0x00, 0x00, 0x00)

#     # ── List Paragraph ────────────────────────────────────────
#     try:
#         list_style = doc.styles['List Paragraph']
#     except KeyError:
#         list_style = doc.styles.add_style('List Paragraph', WD_STYLE_TYPE.PARAGRAPH)
#         list_style.base_style = doc.styles['Normal']
#     list_style.font.name = 'Calibri'
#     list_style.font.size = Pt(11)
#     list_style.paragraph_format.left_indent = Inches(0.5)
#     list_style.paragraph_format.first_line_indent = Inches(-0.25)

#     # ── Code Style ────────────────────────────────────────────
#     try:
#         code_style = doc.styles['Code']
#     except KeyError:
#         code_style = doc.styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
#         code_style.base_style = doc.styles['Normal']
#     code_style.font.name = 'Courier New'
#     code_style.font.size = Pt(9)
#     code_style.paragraph_format.left_indent = Inches(0.25)
#     code_style.paragraph_format.space_before = Pt(6)
#     code_style.paragraph_format.space_after = Pt(6)
#     pPr_code = get_or_add_style_pPr(code_style)
#     shd_code = OxmlElement('w:shd')
#     shd_code.set(qn('w:val'), 'clear')
#     shd_code.set(qn('w:color'), 'auto')
#     shd_code.set(qn('w:fill'), 'F5F5F5')
#     pPr_code.append(shd_code)

#     # ── Caption Style ─────────────────────────────────────────
#     try:
#         caption_style = doc.styles['Caption']
#     except KeyError:
#         caption_style = doc.styles.add_style('Caption', WD_STYLE_TYPE.PARAGRAPH)
#         caption_style.base_style = doc.styles['Normal']
#     caption_style.font.name = 'Calibri'
#     caption_style.font.size = Pt(10)
#     caption_style.font.italic = True
#     caption_style.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
#     caption_style.paragraph_format.space_before = Pt(4)

#     # ── Table Paragraph (pandoc uses this for table cell body text) ──
#     try:
#         tbl_para = doc.styles['Table Paragraph']
#     except KeyError:
#         tbl_para = doc.styles.add_style('Table Paragraph', WD_STYLE_TYPE.PARAGRAPH)
#         tbl_para.base_style = doc.styles['Normal']
#     tbl_para.font.name = 'Calibri'
#     tbl_para.font.size = Pt(10)
#     tbl_para.paragraph_format.space_before = Pt(3)
#     tbl_para.paragraph_format.space_after = Pt(3)
#     tbl_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
#     tbl_para.paragraph_format.line_spacing = 1.0

#     # ── Table Heading (pandoc uses this for header row cells) ────────
#     try:
#         tbl_heading = doc.styles['Table Heading']
#     except KeyError:
#         tbl_heading = doc.styles.add_style('Table Heading', WD_STYLE_TYPE.PARAGRAPH)
#         tbl_heading.base_style = doc.styles['Normal']
#     tbl_heading.font.name = 'Calibri'
#     tbl_heading.font.size = Pt(10)
#     tbl_heading.font.bold = True
#     tbl_heading.paragraph_format.space_before = Pt(3)
#     tbl_heading.paragraph_format.space_after = Pt(3)
#     force_style_rPr_color(tbl_heading, 0xFF, 0xFF, 0xFF)
#     # Dark background on header cells via paragraph shading
#     pPr_th = get_or_add_style_pPr(tbl_heading)
#     shd_th = OxmlElement('w:shd')
#     shd_th.set(qn('w:val'), 'clear')
#     shd_th.set(qn('w:color'), 'auto')
#     shd_th.set(qn('w:fill'), '2E2E2E')
#     pPr_th.append(shd_th)

#     # ── Table Grid style — force visible borders ──────────────
#     patch_table_grid_borders(doc, color='CCCCCC', sz='4')

#     # ── Footer ────────────────────────────────────────────────
#     footer = section.footer
#     footer_para = footer.paragraphs[0]
#     footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

#     fp_pPr = footer_para._p.get_or_add_pPr()
#     fp_pBdr = OxmlElement('w:pBdr')
#     fp_top = OxmlElement('w:top')
#     fp_top.set(qn('w:val'), 'single')
#     fp_top.set(qn('w:sz'), '4')
#     fp_top.set(qn('w:space'), '1')
#     fp_top.set(qn('w:color'), '666666')
#     fp_pBdr.append(fp_top)
#     fp_pPr.append(fp_pBdr)

#     run_label = footer_para.add_run("[Dashboard Name] | Story Guide | Page ")
#     run_label.font.name = 'Calibri'
#     run_label.font.size = Pt(9)
#     run_label.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
#     add_page_number(footer_para)

#     # ── Sample Content (demonstrates every style for pandoc) ─────

#     # H1
#     doc.add_heading('1. Dashboard at a Glance', level=1)

#     # Normal body text
#     doc.add_paragraph(
#         'This reference document shows the styles used in the Risk Management Story Guide.'
#     )

#     # H2
#     doc.add_heading('1.1 What is this dashboard?', level=2)

#     # Bold label + normal run in one paragraph
#     p = doc.add_paragraph(style='Normal')
#     r_bold = p.add_run('TL;DR — ')
#     r_bold.bold = True
#     p.add_run(
#         'A single view of risk adjustment performance tracking RAF capture, '
#         'coding gaps, and recapture rates.'
#     )

#     # H2 before table
#     doc.add_heading('Sample Table', level=2)

#     # ── Table ─────────────────────────────────────────────────
#     table = doc.add_table(rows=4, cols=3)
#     table.style = 'Table Grid'

#     tbl = table._tbl
#     tblPr = tbl.find(qn('w:tblPr'))
#     if tblPr is None:
#         tblPr = OxmlElement('w:tblPr')
#         tbl.insert(0, tblPr)
#     tblW = OxmlElement('w:tblW')
#     tblW.set(qn('w:w'), '5000')
#     tblW.set(qn('w:type'), 'pct')
#     tblPr.append(tblW)

#     headers = ['Table', 'Column', 'Role']
#     for cell, hdr in zip(table.rows[0].cells, headers):
#         set_cell_background(cell, '2E2E2E')
#         set_cell_padding(cell, 6)
#         cell.paragraphs[0].clear()
#         run = cell.paragraphs[0].add_run(hdr)
#         run.font.name = 'Calibri'
#         run.font.size = Pt(10)
#         run.font.bold = True
#         run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

#     row_data = [
#         ('risk_core',   'risk_value',      'HCC risk weight'),
#         ('attribution', 'member_count',    'Member denominator'),
#         ('date',        'month_of_date',   'Time intelligence'),
#     ]
#     for i, (row, data) in enumerate(zip(table.rows[1:], row_data), 1):
#         bg = 'FFFFFF' if i % 2 == 1 else 'F5F5F5'
#         for cell, text in zip(row.cells, data):
#             set_cell_background(cell, bg)
#             set_cell_padding(cell, 6)
#             cell.paragraphs[0].clear()
#             run = cell.paragraphs[0].add_run(text)
#             run.font.name = 'Calibri'
#             run.font.size = Pt(10)

#     add_border(table, color='CCCCCC', sz='4')

#     # H2 before bullets
#     doc.add_heading('Sample Bullets', level=2)

#     # Bullet items (List Paragraph style)
#     doc.add_paragraph('First bullet item', style='List Paragraph')
#     p_bullet = doc.add_paragraph(style='List Paragraph')
#     r_label = p_bullet.add_run('Bold label  ')
#     r_label.bold = True
#     p_bullet.add_run('followed by normal text')

#     # H3 — DAX label
#     doc.add_heading('DAX measure(s):', level=3)

#     # Code block
#     code_para = doc.add_paragraph(style='Code')
#     code_para.add_run(
#         'RAF recapture rate = DIVIDE(SUM(risk_core[risk_value]),\n'
#         'SUM(risk_core[risk_denominator]))'
#     )

#     # Caption / italic note
#     doc.add_paragraph(
#         'Causation note: This metric reflects outcomes, not causes. '
#         'Always investigate the drivers.',
#         style='Caption'
#     )

#     # ── Save ──────────────────────────────────────────────────
#     output_path = str(OUTPUT_ROOT / "reference.docx")
#     OUTPUT_ROOT.mkdir(parents=True, exist_ok=True)
#     doc.save(output_path)
#     print(f"✅ reference.docx created: {output_path}")


# if __name__ == "__main__":
#     create_reference_doc()


import os
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_ROOT = Path(__file__).resolve().parent.parent.parent / "output"


def set_cell_background(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex.lstrip('#'))
    tcPr.append(shd)


def set_cell_padding(cell, padding_pt):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side in ('top', 'left', 'bottom', 'right'):
        margin = OxmlElement(f'w:{side}')
        margin.set(qn('w:w'), str(int(padding_pt * 20)))
        margin.set(qn('w:type'), 'dxa')
        tcMar.append(margin)
    tcPr.append(tcMar)


def add_border(table, color='CCCCCC', sz='4'):
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for edge in ('top', 'left', 'bottom', 'right'):
                el = OxmlElement(f'w:{edge}')
                el.set(qn('w:val'), 'single')
                el.set(qn('w:sz'), sz)
                el.set(qn('w:space'), '0')
                el.set(qn('w:color'), color)
                tcBorders.append(el)
            tcPr.append(tcBorders)


def add_page_number(paragraph):
    run1 = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run1._r.append(fldChar1)

    run2 = paragraph.add_run()
    instrText = OxmlElement('w:instrText')
    instrText.text = 'PAGE'
    instrText.set(qn('xml:space'), 'preserve')
    run2._r.append(instrText)

    run3 = paragraph.add_run()
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run3._r.append(fldChar2)

    for run in (run1, run2, run3):
        run.font.name = 'Calibri'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def get_or_add_style_pPr(style):
    element = style._element
    pPr = element.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        element.append(pPr)
    return pPr


def force_style_rPr_color(style, r, g, b):
    element = style._element
    rPr = element.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        element.append(rPr)
    for existing in rPr.findall(qn('w:color')):
        rPr.remove(existing)
    color_el = OxmlElement('w:color')
    color_el.set(qn('w:val'), f'{r:02X}{g:02X}{b:02X}')
    rPr.append(color_el)
    theme_color = rPr.find(qn('w:themeColor'))
    if theme_color is not None:
        rPr.remove(theme_color)


def patch_table_grid_borders(doc, color='CCCCCC', sz='4'):
    try:
        tg = doc.styles['Table Grid']
    except KeyError:
        return
    tbl_el = tg._element
    tblPr = tbl_el.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl_el.append(tblPr)
    for old in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(old)
    tblBorders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), sz)
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color)
        tblBorders.append(el)
    tblPr.append(tblBorders)


def create_reference_doc():
    doc = Document()

    # ── Page Setup (A4) ───────────────────────────────────────
    section = doc.sections[0]
    section.page_width  = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1)
    section.right_margin  = Inches(1)

    # ── Normal / Body Text  (11pt) ────────────────────────────
    normal = doc.styles['Normal']
    normal.font.name  = 'Calibri'
    normal.font.size  = Pt(11)
    normal.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    normal.paragraph_format.space_after       = Pt(6)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing      = 1.15

    # ── Heading 1  (16pt, bold, black, underline border) ──────
    h1 = doc.styles['Heading 1']
    h1.font.name  = 'Calibri'
    h1.font.size  = Pt(16)
    h1.font.bold  = True
    h1.paragraph_format.space_before   = Pt(24)
    h1.paragraph_format.space_after    = Pt(12)
    h1.paragraph_format.keep_with_next = True
    force_style_rPr_color(h1, 0x00, 0x00, 0x00)
    pPr_h1 = get_or_add_style_pPr(h1)
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr_h1.append(pBdr)

    # ── Heading 2  (12pt, bold, black) ────────────────────────
    h2 = doc.styles['Heading 2']
    h2.font.name   = 'Calibri'
    h2.font.size   = Pt(12)
    h2.font.bold   = True
    h2.font.italic = False
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after  = Pt(6)
    force_style_rPr_color(h2, 0x00, 0x00, 0x00)

    # ── Heading 3  (12pt, bold, black) ────────────────────────
    h3 = doc.styles['Heading 3']
    h3.font.name   = 'Calibri'
    h3.font.size   = Pt(12)
    h3.font.bold   = True
    h3.font.italic = False
    h3.paragraph_format.space_before = Pt(8)
    h3.paragraph_format.space_after  = Pt(4)
    force_style_rPr_color(h3, 0x00, 0x00, 0x00)

    # ── List Paragraph ────────────────────────────────────────
    try:
        list_style = doc.styles['List Paragraph']
    except KeyError:
        list_style = doc.styles.add_style('List Paragraph', WD_STYLE_TYPE.PARAGRAPH)
        list_style.base_style = doc.styles['Normal']
    list_style.font.name = 'Calibri'
    list_style.font.size = Pt(11)
    list_style.paragraph_format.left_indent       = Inches(0.5)
    list_style.paragraph_format.first_line_indent = Inches(-0.25)

    # ── Code  (9pt, Courier New, light grey bg) ───────────────
    try:
        code_style = doc.styles['Code']
    except KeyError:
        code_style = doc.styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
        code_style.base_style = doc.styles['Normal']
    code_style.font.name = 'Courier New'
    code_style.font.size = Pt(9)
    code_style.paragraph_format.left_indent  = Inches(0.25)
    code_style.paragraph_format.space_before = Pt(6)
    code_style.paragraph_format.space_after  = Pt(6)
    pPr_code = get_or_add_style_pPr(code_style)
    shd_code = OxmlElement('w:shd')
    shd_code.set(qn('w:val'),   'clear')
    shd_code.set(qn('w:color'), 'auto')
    shd_code.set(qn('w:fill'),  'F5F5F5')
    pPr_code.append(shd_code)

    # ── Caption ───────────────────────────────────────────────
    try:
        caption_style = doc.styles['Caption']
    except KeyError:
        caption_style = doc.styles.add_style('Caption', WD_STYLE_TYPE.PARAGRAPH)
        caption_style.base_style = doc.styles['Normal']
    caption_style.font.name   = 'Calibri'
    caption_style.font.size   = Pt(10)
    caption_style.font.italic = True
    caption_style.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    caption_style.paragraph_format.space_before = Pt(4)

    # ── Table Paragraph (pandoc body cells) ───────────────────
    try:
        tbl_para = doc.styles['Table Paragraph']
    except KeyError:
        tbl_para = doc.styles.add_style('Table Paragraph', WD_STYLE_TYPE.PARAGRAPH)
        tbl_para.base_style = doc.styles['Normal']
    tbl_para.font.name = 'Calibri'
    tbl_para.font.size = Pt(10)
    tbl_para.paragraph_format.space_before       = Pt(3)
    tbl_para.paragraph_format.space_after        = Pt(3)
    tbl_para.paragraph_format.line_spacing_rule  = WD_LINE_SPACING.MULTIPLE
    tbl_para.paragraph_format.line_spacing       = 1.0

    # ── Table Heading (pandoc header cells) ───────────────────
    try:
        tbl_heading = doc.styles['Table Heading']
    except KeyError:
        tbl_heading = doc.styles.add_style('Table Heading', WD_STYLE_TYPE.PARAGRAPH)
        tbl_heading.base_style = doc.styles['Normal']
    tbl_heading.font.name = 'Calibri'
    tbl_heading.font.size = Pt(10)
    tbl_heading.font.bold = True
    tbl_heading.paragraph_format.space_before = Pt(3)
    tbl_heading.paragraph_format.space_after  = Pt(3)
    force_style_rPr_color(tbl_heading, 0xFF, 0xFF, 0xFF)
    pPr_th = get_or_add_style_pPr(tbl_heading)
    shd_th = OxmlElement('w:shd')
    shd_th.set(qn('w:val'),   'clear')
    shd_th.set(qn('w:color'), 'auto')
    shd_th.set(qn('w:fill'),  '2E2E2E')
    pPr_th.append(shd_th)

    # ── Table Grid — visible borders ──────────────────────────
    patch_table_grid_borders(doc, color='CCCCCC', sz='4')

    # ── Footer  (page number) ─────────────────────────────────
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    fp_pPr = footer_para._p.get_or_add_pPr()
    fp_pBdr = OxmlElement('w:pBdr')
    fp_top = OxmlElement('w:top')
    fp_top.set(qn('w:val'),   'single')
    fp_top.set(qn('w:sz'),    '4')
    fp_top.set(qn('w:space'), '1')
    fp_top.set(qn('w:color'), '666666')
    fp_pBdr.append(fp_top)
    fp_pPr.append(fp_pBdr)

    run_label = footer_para.add_run("[Dashboard Name] | Story Guide | Page ")
    run_label.font.name  = 'Calibri'
    run_label.font.size  = Pt(9)
    run_label.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    add_page_number(footer_para)

    # ── Sample content (so pandoc can see every style) ────────
    doc.add_heading('1. Dashboard at a Glance', level=1)
    doc.add_paragraph('Body text at 11pt Calibri.')

    doc.add_heading('1.1 Sub-heading', level=2)
    p = doc.add_paragraph(style='Normal')
    p.add_run('Bold label — ').bold = True
    p.add_run('followed by normal 11pt text.')

    doc.add_heading('DAX measure(s):', level=3)
    doc.add_paragraph(
        'Documented risk = CALCULATE(DIVIDE(SUM(risk_core[risk_value]),'
        'SUM(risk_core[patient_count])))',
        style='Code'
    )

    doc.add_heading('Sample Table', level=2)
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Table Grid'
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr')) or OxmlElement('w:tblPr')
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '5000')
    tblW.set(qn('w:type'), 'pct')
    tblPr.append(tblW)

    for cell, hdr in zip(table.rows[0].cells, ['Table', 'Column', 'Role']):
        set_cell_background(cell, '2E2E2E')
        set_cell_padding(cell, 6)
        run = cell.paragraphs[0].add_run(hdr)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for i, row in enumerate(table.rows[1:], 1):
        bg = 'FFFFFF' if i % 2 == 1 else 'F5F5F5'
        for cell, text in zip(row.cells, ['risk_core', 'risk_value', 'HCC risk weight']):
            set_cell_background(cell, bg)
            set_cell_padding(cell, 6)
            cell.paragraphs[0].add_run(text)

    add_border(table)

    # ── Save ──────────────────────────────────────────────────
    output_path = str(OUTPUT_ROOT / "reference.docx")
    OUTPUT_ROOT.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f"✅ reference.docx saved: {output_path}")


if __name__ == "__main__":
    create_reference_doc()